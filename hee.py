from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Allianz PNB Life Voice Bot Script', 0)

# First Section - Calls landline/mobile number on file
doc.add_heading('Calls landline/mobile number on file', level=1)
doc.add_paragraph("Bot Initiates:\n[AI]: Greetings from Allianz PNB Life. May I speak to (customer’s name), please?")

# Scenario 1: Speaking to Policy Owner
doc.add_heading('Scenario 1: Speaking to Policy Owner', level=1)
doc.add_paragraph("User Response:\nYes, speaking (or any other response of the same meaning).")

# Scenario 2: Speaking to Someone Else
doc.add_heading('Scenario 2: Speaking to Someone Else', level=1)
doc.add_paragraph("User Response:\nHe/She is not available at the moment (or any other response of the same meaning).")
doc.add_paragraph("Bot:\n[AI]: Thank you for that information. We will call again some other time.\n<End call>")

# New Section: Policy Number Required
doc.add_heading('New Section: Policy Number Required', level=1)
doc.add_paragraph("User: I don’t know the policy holder’s name.")
doc.add_paragraph("Bot:\n[AI]: Please provide your policy number to fetch your details and update the name on the record.")

# Policy Number Provided
doc.add_heading('Policy Number Provided', level=1)
doc.add_paragraph("User: 12345678")
doc.add_paragraph("Bot:\n[AI]: Thank you! Let me verify this information.")

# If Policy Number Matches
doc.add_heading('If Policy Number Matches', level=1)
doc.add_paragraph("Bot:\n[AI]: The policy number matches our records. Is this the right time to talk?")
doc.add_paragraph("User: No, he/she is busy.")
doc.add_paragraph("Bot:\n[AI]: Thank you for informing us. Allianz will call you again another time.\n<End call>")

# If Policy Number Does Not Match
doc.add_heading('If Policy Number Does Not Match', level=1)
doc.add_paragraph("Bot:\n[AI]: The policy number you provided does not match our records. Please contact our office to rectify the issue.\n<End call>")

# Before Due Date Reminder
doc.add_heading('Before Due Date Reminder', level=1)
doc.add_paragraph("Bot:\n[AI]: Hi (customer’s name). Please be aware that this call may be recorded for security and quality assurance purposes. "
    "We wish to remind you of your premium payment for your (plan name) policy with policy number ending in (last 4 digits).\n"
    "May we ask you to kindly pay your premium of (premium amount) (currency) on or before (due date) to keep your policy active and enjoy continuous coverage.")

# After Due Date Reminder
doc.add_heading('After Due Date Reminder', level=1)
doc.add_paragraph("Bot:\n[AI]: Hi (customer’s name). Please be aware that this call may be recorded for security and quality assurance purposes. "
    "We wish to remind you of your premium payment for your (plan name) policy with number ending in (last 4 digits).\n"
    "You have missed your payment of (premium amount) (currency) which was due last (due date).\n"
    "Keep your policy active and enjoy continuous coverage. May we ask you to pay your premium on or before (31 days after due date – actual date)?")

# Customer Acknowledges Payment Delay
doc.add_heading('Customer Acknowledges Payment Delay', level=1)
doc.add_paragraph("User: Yes, I have yet to make the payment (or similar responses).")
doc.add_paragraph("Bot:\n[AI]: You may pay your premium over the counter through our partner banks and payment centers.\n"
    "For Philippine Peso and US Dollar policies, you may pay via PNB, Metrobank, and BDO.\n"
    "For Philippine Peso policies, you may also pay via Cebuana Lhuillier.\n"
    "You also have the option to pay your premium online. For Philippine Peso and US Dollar policies, you may pay via BDO Online, "
    "PNB Mobile Banking, PNB Internet Banking, and Metrobank Direct Online. Please note, however, that Metrobank Mobile App, "
    "GCash, Maya, and Bancnet Online are only applicable to Philippine Peso policies.\n"
    "You may also log on to your Allianz Touch account and pay your premium via the PayNow option.\n"
    "While our agents and financial advisors are happy to help you, they aren’t authorized to receive payments. Please pay your premiums "
    "directly to our official payment facilities to ensure timely recording.")

# Customer Informs Payment Has Been Made
doc.add_heading('Customer Informs Payment Has Been Made', level=1)
doc.add_paragraph("User:\nI have already paid this on ___.\nOR\nMy policy is under auto-pay. This should have been paid.\nOR\n(similar responses).")
doc.add_paragraph("Bot:\n[AI]: Thank you for that information. Kindly disregard this reminder if payment has already been made. "
    "For policies enrolled under auto-pay, an automatic rebilling will be made in the following days. Kindly ensure your account is active and well-funded.")

# Customer Asks About Allianz Touch
doc.add_heading('Customer Asks About Allianz Touch', level=1)
doc.add_paragraph("User:\nHow do I register to/log on to Allianz Touch? (or similar responses).")
doc.add_paragraph("Bot:\n[AI]: Simply go to www.touch.allianzpnblife.ph and use the email address you provided during your policy application when creating/logging-in to your account.")

# Customer Asks for Information Again
doc.add_heading('Customer Asks for Information Again', level=1)
doc.add_paragraph("User:\nCan you repeat the link/information? (or similar responses).")
doc.add_paragraph("Bot:\n[AI]: It’s www.touch.allianzpnblife.ph. Reminder to use the email address you provided during your policy application when creating/logging-in to your account.")

# Customer Agrees to Pay Premium
doc.add_heading('Customer Agrees to Pay Premium', level=1)
doc.add_paragraph("User:\nThanks for informing me. Will pay my premium accordingly (or any similar response).")
doc.add_paragraph("Bot:\n[AI]: My pleasure speaking with you, Mr./Ms. (Name). For other concerns, feel free to reach out to us via email at customercare@allianzpnblife.ph or call us at 8818-4357.\n"
    "Once again, thank you for choosing Allianz PNB Life as your insurance partner. Have a good day ahead!\n<End call>")

# Save the document
file_path = '/mnt/data/Allianz_PNB_Life_Voice_Bot_Script_Updated.docx'
doc.save(file_path)

file_path
